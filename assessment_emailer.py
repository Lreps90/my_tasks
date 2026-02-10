import smtplib
import os
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document

# Set your email and password
sender_email = "lreps@heathside.surrey.sch.uk"
password = ""  # Remember to handle this securely
smtp_server = "smtp.office365.com"  # Adjust for your email provider
smtp_port = 587  # SMTP port for your email provider

# Path to the directory containing the Word documents
documents_directory = r"C:\Users\Lreps\OneDrive - Heathside School\Computer Science\GCSE\10S\System Architecture Responses - Copy"


def extract_email(doc):
    first_paragraph_text = doc.paragraphs[0].text.strip().lower()
    # Check if the document title contains "%_Feedback"
    if "%_feedback" in first_paragraph_text:
        # If so, the email address is expected to be in the second line
        if len(doc.paragraphs) > 1:
            return doc.paragraphs[1].text.strip()
    else:
        # Otherwise, extract email from the first line
        if "email:" in first_paragraph_text:
            return first_paragraph_text.split("email:")[1].strip()
        return first_paragraph_text
    return None

def send_email_with_attachment(recipient_email, document_path):
    try:
        # Create a multipart message
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = recipient_email
        message['Subject'] = "Systems Architecture Assessment Documents"

        # Add your custom message here
        custom_message = """Hello,

Attached you will find one of two documents you should recieve relating to your Systems Architecture assessment. One is your original questions and the other is a feedback document containing your marks.

Regards,
Mr Reps."""
        message.attach(MIMEText(custom_message, "plain"))

        # Attach the Word document
        with open(document_path, "rb") as attachment:
            part = MIMEApplication(attachment.read(), Name=os.path.basename(document_path))
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(document_path)}"'
        message.attach(part)

        # Log in to server and send the email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, password)
        server.sendmail(sender_email, recipient_email, message.as_string())
        server.quit()
        print(f"Email sent to {recipient_email} with attachment {document_path}")
    except Exception as e:
        print(f"Failed to send email to {recipient_email}. Error: {e}")

# Iterate over each document in the folder
for filename in os.listdir(documents_directory):
    if filename.endswith(".docx"):  # Make sure to process .docx files only
        document_path = os.path.join(documents_directory, filename)
        doc = Document(document_path)
        recipient_email = extract_email(doc)  # Determine the recipient email address
        if recipient_email:
            send_email_with_attachment(recipient_email, document_path)
        else:
            print(f"No valid email found in document: {filename}")
