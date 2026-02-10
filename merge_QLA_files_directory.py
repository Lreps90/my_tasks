import os
from docx import Document

def open_and_close_documents(folder_path):
    """
    Open and close each Word document in the folder to ensure they are accessible.

    :param folder_path: Path to the folder containing Word documents.
    """
    doc_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    doc_files.sort()  # Optional: Sort files alphabetically

    print(f"Checking {len(doc_files)} Word documents by opening and closing them.")

    for doc_file in doc_files:
        file_path = os.path.join(folder_path, doc_file)

        try:
            print(f"Opening file: {file_path}")
            document = Document(file_path)  # Open the document
            document.save(file_path)  # Save the document to ensure it can be written to
            print(f"Successfully opened and closed {file_path}.")
        except Exception as e:
            print(f"Error opening/closing file {file_path}: {e}")
            raise

def merge_word_documents(folder_path, output_filename):
    """
    Merge all Word documents in the specified folder into one document consecutively.

    :param folder_path: Path to the folder containing Word documents.
    :param output_filename: Name of the output merged document.
    """
    # Get all .docx files in the folder
    doc_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    doc_files.sort()  # Optional: Sort files alphabetically

    print(f"Found {len(doc_files)} Word documents to merge.")

    # Create a new Word document
    merged_document = Document()

    for doc_file in doc_files:
        file_path = os.path.join(folder_path, doc_file)

        try:
            print(f"Processing file: {file_path}")

            # Open the Word document
            current_document = Document(file_path)
            print(f"Successfully opened {file_path}.")

            # Append each paragraph from the current document
            for paragraph in current_document.paragraphs:
                merged_document.add_paragraph(paragraph.text)

            # Add a page break between documents
            merged_document.add_page_break()

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            raise

    # Save the merged document
    merged_document.save(output_filename)
    print(f"All documents have been merged into {output_filename}")


# Example usage
if __name__ == "__main__":
    folder_path = r"C:\Users\Lreps\OneDrive - Heathside School\Computer Science\GCSE\Assessments\2024-2025\QLA"
    output_filename = r"C:\Users\Lreps\OneDrive - Heathside School\Computer Science\GCSE\Assessments\2024-2025\QLA\All_QLA.docx"

    # Step 1: Open and close each document
    open_and_close_documents(folder_path)

    # Step 2: Merge the documents
    merge_word_documents(folder_path, output_filename)
