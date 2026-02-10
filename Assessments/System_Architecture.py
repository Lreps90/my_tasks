import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import os
from docx import Document
#
# # This program formats the Unit 1 - System Architecture for AI Marking.
# # AI marking required for all questions
# #
# # When uploading a spreadsheet the first question should on column E.
#
# # CHATGPT PROMTP:
#
# Task: Feedback on students assessment
# 1.	Feedback for Each Question:
# •	Act as a GCSE Computer science teacher and for Q1 to Q17, provide very detailed feedback and allocate marks according to your instructions and the mark scheme. You may award half marks for questions that lack clarity and depth.
# •	Refer to the ‘System Architecture Mark Scheme chatGPT’ document
# •	Questions that are vague or lack clarity or depth should not be awarded marks but you may use your judgement in awarding half marks.
# •	Do not add up marks at this stage. Always do this in python.
# 2.	Total Marks and Percentage:
# •	Use python to tally and count up the total marks awarded
# •	Use python to calculate the total marks out of 40.
# •	Compute the percentage score and round it to the nearest whole number
# 3.	Document Creation:
# •	Use python to create a downloadable microsoft word document that contains the marks and feedback for each question.
# •	Title: "StudentName_PercentageScore_Feedback", where StudentName is the student name and PercentageScore is the students percentage score
# •	Begin with the student's email.
# •	Include detailed feedback for each of the 17 questions.
# •	Conclude with the overall summary, total marks, and percentage score.
# •	Put the students email address on the top line. This should be found in the students assessment
# 4.	Save and Provide Download Link:
# •	Save the document and generate a link for downloading.
# 5.	Extras:
# •	Do not stop at any point during our response.



# This program formats the Unit 1 - System Architecture for AI Marking.
# AI marking required for all questions
#
# When uploading a spreadsheet the first question should on column E.

def select_file():
    global input_file_path
    input_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if input_file_path:
        file_label.config(text=os.path.basename(input_file_path))

def select_output_folder():
    global output_folder_path
    output_folder_path = filedialog.askdirectory()
    if output_folder_path:
        folder_label.config(text=output_folder_path)

def split_excel_to_word():
    try:
        # Read the Excel file
        df = pd.read_excel(input_file_path)

        # Filter out columns starting with 'Feedback' or 'Points'
        df = df.loc[:, ~df.columns.str.startswith('Feedback')]
        df = df.loc[:, ~df.columns.str.startswith('Points')]

        # Iterate through each row to create a Word document for each student
        for index, row in df.iterrows():
            doc = Document()

            # Use the 'Name' column to create a filename for the Word document
            student_name = row['Name']
            filename = os.path.join(output_folder_path, f"{student_name}.docx")

            # Add the student's Email and Name at the beginning of the document
            doc.add_paragraph(f"Email: {row['Email']}")
            doc.add_paragraph(f"Name: {row['Name']}")
            doc.add_paragraph("")  # Add a space after the student's name

            # Counter for numbering the questions
            question_number = 1

            # Iterate through columns/questions, skipping 'Email' and 'Name'
            for question in df.columns[2:]:  #
                # Format with question number, then question, and student's response on a new line
                question_text = f"Question {question_number}: {question}\nAnswer: {row[question]}"
                doc.add_paragraph(question_text)
                doc.add_paragraph("")  # Add a space between entries for readability

                question_number += 1  # Increment the question number

            # Save the document
            doc.save(filename)

        messagebox.showinfo("Success", "All documents have been created successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# GUI setup
root = tk.Tk()
root.title("Excel to Word Splitter")

input_file_path = ''
output_folder_path = ''

tk.Button(root, text="Select Excel File", command=select_file).pack()
file_label = tk.Label(root, text="No file selected")
file_label.pack()

tk.Button(root, text="Select Output Folder", command=select_output_folder).pack()
folder_label = tk.Label(root, text="No folder selected")
folder_label.pack()

tk.Button(root, text="Process", command=split_excel_to_word).pack()

root.mainloop()


