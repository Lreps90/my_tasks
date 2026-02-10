import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import os
from docx import Document

# This program formats the Unit 8 - Logic and language Assessment for AI Marking.
# AI marking required for questions 12, 30, 31, 32 and 33.
# Manual marking required for question 13 (add marks to spreadsheet manually).
# When uploading a spreadsheet the first question should on column E.

# CHATGPT PROMTP:

# Title: Automated Assessment and Feedback Generation with Detailed Mark Allocation for Student Work
#
# Objective: To evaluate student responses against a predefined mark scheme, provide detailed feedback including marks awarded for each question, calculate total marks, and generate an overall performance summary. This process encompasses both manual evaluation of specific questions and automated tallying of total scores.
#
# Tasks:
#
# Document Evaluation:
#
# Concentrate on evaluating student responses within a specified document.
# Identify sections requiring manual marking and evaluate them against the provided mark scheme, taking into account accuracy, completeness, and coherence.
# For each manually marked question, provide feedback and explicitly state the marks awarded out of the total available for that question.
# Mark Allocation:
#
# Use subjective judgment based on the criteria to allocate points for manual marking. Ensure fairness and justify the marks awarded.
# Include automatically awarded points as indicated within the document in the total score calculation.
# Feedback Generation:
#
# Offer individual feedback for each question marked manually, highlighting strengths and areas for improvement.
# Include the mark awarded for each question, indicating both the points achieved and the points available.
# Provide constructive suggestions to enhance the student's understanding and performance.
# Total Score Calculation:
#
# Tally both manually and automatically awarded points to calculate the total score.
# Convert the total score out of the maximum available points into a percentage for a comprehensive assessment. The total maximum points available is 51.
# Performance Summary:
#
# Summarise the student's performance, noting particular areas of strength and recommending areas for further development.
# Document Creation:
#
# Use Python to compile the feedback, marks awarded, and performance summary into a Microsoft Word document.
# The document should detail the feedback and marks awarded for each question individually and provide a final mark as a percentage. You should justify and provide detailed feedback as to why you have awarded each mark per question. Name the document "studentName_FEEDBACK_percentage," where studentName is the student's name and percentage is the percentage score awarded to the nearest whole number.
# Deliverables:
#
# A detailed feedback document for each evaluated section, including marks awarded for each question and giving the overall mark out of 51 and percentage awarded.
# A summary of the student's overall performance with constructive feedback.
# A downloadable Microsoft Word document containing all feedback and the final mark as a percentage.

def select_file():
    global input_file_path
    input_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if input_file_path:
        file_label.config(text=os.path.basename(input_file_path))

def select_output_folder():
    global output_folder_path
    output_folder_path = filedialog.askdirectory()
    if output_folder_path:
        folder_label.config(text=output_folder_path)

def split_excel_to_word():
    if not input_file_path or not output_folder_path:
        messagebox.showerror("Error", "Please select both an Excel file and an output folder.")
        return

    df = pd.read_excel(input_file_path, engine='openpyxl')

    # Remove columns that start with "Feedback"
    df = df.loc[:, ~df.columns.str.startswith('Feedback')]

    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    for index, row in df.iterrows():
        total_points = 0
        name = str(row['Name']).strip().replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
        doc = Document()  # Create a new Document
        doc.add_heading('Responses for ' + name, 0)  # Add a heading with the name

        question_number = 1  # Initialize question number
        for i in range(4, len(df.columns), 2):  # Assuming questions start from column E
            response = str(row[df.columns[i]]).strip()

            if question_number == 10 and "length" in response.lower():
                points_awarded = 1
            elif question_number == 11 and ("format" in response.lower() or "range" in response.lower()):
                points_awarded = 1
            elif question_number in [12, 30, 31, 32, 33]:
                points_awarded = "Requires manual marking"
            elif not pd.isna(row[df.columns[i + 1]]):
                points_awarded = int(str(row[df.columns[i + 1]]).strip())
            else:
                points_awarded = 0

            if isinstance(points_awarded, int):
                total_points += points_awarded

            if question_number == 12:
                # Custom algorithm text for question 12
                algorithm_text = ("The following is an incomplete algorithm for checking that a user has entered the correct code. "
                                  "The code entered is to be checked against the code “ABCdef123”. The user is given 3 attempts to get the "
                                  "code correct before being logged out.\n\n"
                                  "1. codeOK = False\n"
                                  "2. attemptsAllowed = 3\n"
                                  "3. attempts = 0\n"
                                  "4. while attempts <= attemptsAllowed and NOT codeOK\n"
                                  "5.\n"
                                  "6.\n"
                                  "7.\n"
                                  "8.\n"
                                  "9.\n"
                                  "10.\n"
                                  "11.\n"
                                  "12.\n"
                                  "13. if codeOK then\n"
                                  "14.       print(\"Welcome back\")\n"
                                  "15. else  \n"
                                  "16.       print(\"Logged out\")\n"
                                  "17. endif\n\n"
                                  "Complete the algorithm that should occupy lines 5-12 (not all lines need to be used but you should not need more). Note 'code' is used to avoid Microsoft thinking I am asking for sensitive information! (5 marks)")
                doc.add_paragraph(f"Question {question_number}:\n{algorithm_text}\nResponse: {response}\nPoints Awarded: {points_awarded}\n")
            elif question_number == 31:
                # New pseudocode text for question 31
                new_question_text = ("The pseudocode below shows an algorithm that is supposed to calculate the stock levels in a bookstore. "
                                     "If the number of books for sale goes below 100, a message is displayed telling the shopkeeper to order more books. "
                                     "Books in stock is updated after a new batch of books arrive.\n\n"
                                     "1.  START\n"
                                     "2.  booksInStock = 300\n"
                                     "3.  booksSoldToday = 50\n"
                                     "4.  IF booksInStock < 100 THEN\n"
                                     "5.      PRINT \"Order More Books\"\n"
                                     "6.  ELSE\n"
                                     "7.      PRINT \"No Need to Order Books\"\n"
                                     "8.  ENDIF\n"
                                     "9.  newBooksArrived = 150\n"
                                     "10. booksInStock = booksInStock - booksSoldToday\n"
                                     "11. PRINT \"Total Books in Stock Now: \", booksInStock\n"
                                     "12. END\n\n"
                                     "Identify the two lines with logic errors in the above pseudocode, and explain why they are logical errors. (4 marks)")
                doc.add_paragraph(f"Question {question_number}:\n{new_question_text}\nResponse: {response}\nPoints Awarded: {points_awarded}\n")
            elif question_number == 32:
                # New pseudocode text for question 32
                new_question_text = ("Programmers usually use an IDE (Integrated Development Environment) "
                                     "to develop their programming code. Some features offered by IDEs include"
                                     " breakpoints, stepping through code and watching variables. Describe "
                                     "how an IDE can be used to debug programs that have logical errors in them. (3 marks)")
                doc.add_paragraph(f"Question {question_number}:\n{new_question_text}\nResponse: {response}\nPoints Awarded: {points_awarded}\n")
            else:
                doc.add_paragraph(f"Question {question_number}: {df.columns[i]}\nResponse: {response}\nPoints Awarded: {points_awarded}\n")

            question_number += 1
        doc.add_paragraph(f"Total Automatically assigned points: {total_points}")
        filename = os.path.join(output_folder_path, f"{name}.docx")
        doc.save(filename)  # Save the document

    messagebox.showinfo("Success", "Excel file has been successfully split into Word documents with custom content for specific questions.")

# GUI setup
root = tk.Tk()
root.title("Excel to Word Splitter")

input_file_path = ''
output_folder_path = ''

tk.Button(root, text="Select Excel File", command=select_file).pack()
file_label = tk.Label(root, text="No file selected")
file_label.pack()

tk.Button(root, text="Select Output Folder", command=select_output_folder).pack()
folder_label = tk.Label(root, text="No folder selected")
folder_label.pack()

tk.Button(root, text="Process", command=split_excel_to_word).pack()

root.mainloop()


